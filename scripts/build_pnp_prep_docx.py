"""Write output/pnp-prep-instructions.docx (human PnP prep)."""
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "pnp-prep-instructions.docx"


def _set_run_font(run, *, bold=False, size=12):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rFonts.set(qn("w:cs"), "Arial")


def add_heading(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_run_font(run, bold=True, size=14)


def add_body(doc, text: str, *, bold=False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    _set_run_font(run, bold=bold, size=12)


def add_list(doc, items: list[str], *, numbered: bool) -> None:
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        _set_run_font(run, size=12)


def build() -> Path:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    add_heading(doc, "Свет мой зеркальце – Инструкция по подготовке прототипа")

    add_heading(doc, "Печать")
    add_body(doc, "Распечатайте пять файлов:")
    add_list(
        doc,
        [
            "Правила: svet-moy-zerkalce-rules.pdf. Формат A4, книжная ориентация. Цвет не обязателен. Обычная офисная бумага. Не режьте этот файл.",
            "Карты запросов (лицо): svet-moy-zerkalce-cards.pdf. Формат A4, альбомная ориентация. Цветная печать желательна (синие рамки). Карты – на плотной бумаге. Размер карты 88×63 мм.",
            "Карты запросов (оборот): svet-moy-zerkalce-card-backs.pdf. Формат A4, альбомная ориентация. Тот же лист, что и лица запросов. Карты – на плотной бумаге.",
            "Карты «Ты восхитителен»: svet-moy-zerkalce-votes.pdf. Формат A4, альбомная ориентация. Печатайте в цвете и на плотной бумаге либо вложите карты в протекторы – карты кидают по столу, обычная офисная бумага мнётся и рвётся.",
            "Карты букв: svet-moy-zerkalce-letters.pdf. Можно взять карты букв из другой игры с буквами либо распечатать из этого файла. Формат A4, книжная ориентация. Цвет не обязателен (чёрный текст). Карты – на плотной бумаге.",
        ],
        numbered=True,
    )

    add_heading(doc, "Рубашки / дуплекс")
    add_body(
        doc,
        "Только карты запросов двусторонние. Лица – svet-moy-zerkalce-cards.pdf, обороты – svet-moy-zerkalce-card-backs.pdf.",
    )
    add_body(
        doc,
        "В svet-moy-zerkalce-cards.pdf несколько страниц – все «Карты запросов · ЛИЦО». В svet-moy-zerkalce-card-backs.pdf столько же страниц – «Карты запросов · ОБОРОТ». Напечатайте страницу N лиц и страницу N оборотов на двух сторонах одного листа. Переворот выберите так, чтобы верх листа остался верхом: после печати текст на обороте читается без поворота страницы. Если лица и рубашки не совпали – смените сторону переворота в настройках принтера.",
    )
    add_body(
        doc,
        "На обороте запросов линий реза нет – это компенсирует возможный рассинхрон принтера при двусторонней печати.",
    )
    add_body(
        doc,
        "svet-moy-zerkalce-votes.pdf печатайте односторонне. Все листы svet-moy-zerkalce-letters.pdf тоже односторонние: оборот бумаги оставьте пустым. Карта буквы, сыгранная рубашкой вверх, – джокер.",
    )

    add_heading(doc, "Резка")
    add_body(
        doc,
        "Разрежьте карты по светло-серым линиям реза (границы карт). Подписи листов («ЛИЦО», «ОБОРОТ», «Ты восхитителен», «Карты букв») отрежьте и выбросьте.",
    )

    add_heading(doc, "Контроль количества")
    add_body(doc, "Должно получиться:")
    add_list(
        doc,
        [
            "14 двусторонних карт запросов;",
            "72 односторонние карты букв, если печатаете из svet-moy-zerkalce-letters.pdf;",
            "6 односторонних карт «Ты восхитителен» (по 1 каждого цвета).",
        ],
        numbered=False,
    )
    add_body(
        doc,
        "Лишние карты «Ты восхитителен» в партии не нужны: при 3–5 игроках уберите лишние цвета в коробку.",
    )

    add_heading(doc, "Вне PnP")
    add_body(
        doc,
        "Карты букв можно взять из другой игры с буквами либо распечатать из svet-moy-zerkalce-letters.pdf. Для партии нужна поверхность стола, достаточная чтобы раскидать карты букв в центре и кидать карты «Ты восхитителен» к игрокам.",
    )

    add_heading(doc, "Контакты")
    add_body(
        doc,
        "При возникновении сложностей: t.me/tanion, vk.ru/tanion, yury@nxt.ru, +79213189331 (не стесняйтесь звонить / писать при возникновении вопросов).",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path}")
